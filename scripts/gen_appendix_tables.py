#!/usr/bin/env python3
"""Generate all appendix tables for the thesis."""

import geopandas as gpd
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DATA_DIR = "/Users/hbox/Documents/GitHub/毕业/data"
OUT_PATH = "/Users/hbox/Documents/GitHub/毕业/output/附录.docx"


def set_cell_font(cell, text, font_name="宋体", font_size=9, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        tblPr.remove(borders)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "0",
            qn("w:color"): "000000",
        })
        borders.append(el)
    tblPr.append(borders)


def make_three_line_table(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else tbl._add_tblPr()
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        tblPr.remove(borders)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "bottom"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "single", qn("w:sz"): "12",
            qn("w:space"): "0", qn("w:color"): "000000",
        })
        borders.append(el)
    for edge in ("bottom",):
        pass
    el = borders.makeelement(qn("w:insideH"), {
        qn("w:val"): "single", qn("w:sz"): "4",
        qn("w:space"): "0", qn("w:color"): "000000",
    })
    borders.append(el)
    el = borders.makeelement(qn("w:insideV"), {
        qn("w:val"): "none", qn("w:sz"): "0",
        qn("w:space"): "0", qn("w:color"): "000000",
    })
    borders.append(el)
    for edge in ("left", "right"):
        el = borders.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "none", qn("w:sz"): "0",
            qn("w:space"): "0", qn("w:color"): "000000",
        })
        borders.append(el)
    tblPr.append(borders)


doc = Document()

style = doc.styles["Normal"]
style.font.name = "宋体"
style.font.size = Pt(10.5)
style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

# ============================================================
# 附录A：省州综合指数均值统计
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("附录A  省州综合指数均值统计")
run.font.name = "黑体"
run.font.size = Pt(14)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run("表A-1  各省州SDGs空间支撑条件综合指数均值（2010—2024）")
run.font.name = "宋体"
run.font.size = Pt(9)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

gdf = gpd.read_file(f"{DATA_DIR}/grid_admin_grouped.gpkg", ignore_geometry=True)
id_col = gdf.columns[0]
grouped = gdf.groupby(["country_group", "admin1_name"])
admin_result = grouped.agg(
    grid_count=(id_col, "count"),
    sdg_2010=("sdg_index_2010", "mean"),
    sdg_2015=("sdg_index_2015", "mean"),
    sdg_2020=("sdg_index_2020", "mean"),
    sdg_2024=("sdg_index_2024", "mean"),
).reset_index()
admin_result["label"] = admin_result["country_group"] + " - " + admin_result["admin1_name"]
admin_result = admin_result.sort_values("sdg_2024", ascending=False)
admin_result = admin_result[~admin_result["country_group"].str.contains("Unmatched")]

table_a = doc.add_table(rows=len(admin_result) + 1, cols=6)
table_a.alignment = WD_TABLE_ALIGNMENT.CENTER
make_three_line_table(table_a)

headers_a = ["省州", "网格数", "2010", "2015", "2020", "2024"]
for i, h in enumerate(headers_a):
    set_cell_font(table_a.rows[0].cells[i], h, bold=True)

for idx, (_, row) in enumerate(admin_result.iterrows()):
    r = idx + 1
    set_cell_font(table_a.rows[r].cells[0], row["label"], alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(table_a.rows[r].cells[1], str(int(row["grid_count"])))
    set_cell_font(table_a.rows[r].cells[2], f"{row['sdg_2010']:.4f}")
    set_cell_font(table_a.rows[r].cells[3], f"{row['sdg_2015']:.4f}")
    set_cell_font(table_a.rows[r].cells[4], f"{row['sdg_2020']:.4f}")
    set_cell_font(table_a.rows[r].cells[5], f"{row['sdg_2024']:.4f}")

p = doc.add_paragraph()
run = p.add_run("注：按2024年综合指数均值降序排列。网格数为10 km网格数量。")
run.font.name = "宋体"
run.font.size = Pt(8)
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

doc.add_page_break()

# ============================================================
# 附录B：熵权零值处理稳健性检验
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("附录B  熵权零值处理稳健性检验")
run.font.name = "黑体"
run.font.size = Pt(14)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

p = doc.add_paragraph()
run = p.add_run("在熵权法计算中，标准化值为0的网格不参与信息熵计算（0×ln(0)→0），可能导致零值占比高的指标权重偏大。本文在标准化值进入熵权计算前增加+1e-6平移处理，确保所有值严格大于0。下表对比平移前后结果。")
run.font.name = "宋体"
run.font.size = Pt(10.5)
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

p = doc.add_paragraph()
run = p.add_run("表B-1  综合指数对比（+1e-6平移前后）")
run.font.name = "宋体"
run.font.size = Pt(9)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

robustness_data = [
    ["2010", "0.418625", "0.418625", "0.000000", "1.000000", "1.000000"],
    ["2015", "0.421544", "0.421544", "0.000000", "1.000000", "1.000000"],
    ["2020", "0.422878", "0.422878", "0.000000", "1.000000", "1.000000"],
    ["2024", "0.424872", "0.424873", "0.000001", "1.000000", "1.000000"],
]

table_b1 = doc.add_table(rows=5, cols=6)
table_b1.alignment = WD_TABLE_ALIGNMENT.CENTER
make_three_line_table(table_b1)
headers_b1 = ["年份", "旧均值", "新均值", "差值", "Pearson r", "Spearman ρ"]
for i, h in enumerate(headers_b1):
    set_cell_font(table_b1.rows[0].cells[i], h, bold=True)
for idx, row_data in enumerate(robustness_data):
    for j, val in enumerate(row_data):
        set_cell_font(table_b1.rows[idx + 1].cells[j], val)

p = doc.add_paragraph()
run = p.add_run("注：Pearson r和Spearman ρ均为1.000000，最大绝对差值<0.000001。+1e-6平移对综合指数结果无实质影响。")
run.font.name = "宋体"
run.font.size = Pt(8)
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

doc.add_page_break()

# ============================================================
# 附录C：极值剪裁范围
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("附录C  极值剪裁范围")
run.font.name = "黑体"
run.font.size = Pt(14)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

p = doc.add_paragraph()
run = p.add_run("表C-1  各指标1%和99%分位数剪裁范围")
run.font.name = "宋体"
run.font.size = Pt(9)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

clip_df = pd.read_csv(f"{DATA_DIR}/outlier_clip_ranges.csv")
table_c = doc.add_table(rows=len(clip_df) + 1, cols=5)
table_c.alignment = WD_TABLE_ALIGNMENT.CENTER
make_three_line_table(table_c)
headers_c = ["指标", "方向", "1%分位数(p01)", "99%分位数(p99)", "数据类型"]
for i, h in enumerate(headers_c):
    set_cell_font(table_c.rows[0].cells[i], h, bold=True)

direction_map = {
    "pop_density": "正向", "nightlight": "正向", "cropland_ratio": "正向",
    "dist_port": "负向", "dist_city": "负向", "dist_major_road": "负向",
    "ndvi": "正向", "forest_ratio": "正向", "npp": "正向",
    "slope": "负向", "climate_water_stress": "负向", "cold_stress": "负向",
}
for idx, (_, row) in enumerate(clip_df.iterrows()):
    r = idx + 1
    ind = row["indicator"]
    set_cell_font(table_c.rows[r].cells[0], ind, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(table_c.rows[r].cells[1], direction_map.get(ind, ""))
    set_cell_font(table_c.rows[r].cells[2], f"{row['p01']:.4f}")
    set_cell_font(table_c.rows[r].cells[3], f"{row['p99']:.4f}")
    set_cell_font(table_c.rows[r].cells[4], "比例" if ind in ["cropland_ratio", "forest_ratio"] else "连续")

doc.add_page_break()

# ============================================================
# 附录D：标准化极值范围
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("附录D  全时期合并标准化极值范围")
run.font.name = "黑体"
run.font.size = Pt(14)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

p = doc.add_paragraph()
run = p.add_run("表D-1  各指标全时期合并极差标准化的全局最小值与最大值")
run.font.name = "宋体"
run.font.size = Pt(9)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

norm_df = pd.read_csv(f"{DATA_DIR}/normalization_ranges.csv")
table_d = doc.add_table(rows=len(norm_df) + 1, cols=4)
table_d.alignment = WD_TABLE_ALIGNMENT.CENTER
make_three_line_table(table_d)
headers_d = ["指标", "全局最小值", "全局最大值", "极差"]
for i, h in enumerate(headers_d):
    set_cell_font(table_d.rows[0].cells[i], h, bold=True)
for idx, (_, row) in enumerate(norm_df.iterrows()):
    r = idx + 1
    set_cell_font(table_d.rows[r].cells[0], row["indicator"], alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(table_d.rows[r].cells[1], f"{row['global_min']:.4f}")
    set_cell_font(table_d.rows[r].cells[2], f"{row['global_max']:.4f}")
    set_cell_font(table_d.rows[r].cells[3], f"{row['global_max'] - row['global_min']:.4f}")

doc.add_page_break()

# ============================================================
# 附录E：固定权重表
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("附录E  固定权重表")
run.font.name = "黑体"
run.font.size = Pt(14)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

p = doc.add_paragraph()
run = p.add_run("表E-1  各指标维度内熵权与最终权重")
run.font.name = "宋体"
run.font.size = Pt(9)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

weight_df = pd.read_csv(f"{DATA_DIR}/weights_multiyear.csv")
table_e = doc.add_table(rows=len(weight_df) + 1, cols=5)
table_e.alignment = WD_TABLE_ALIGNMENT.CENTER
make_three_line_table(table_e)
headers_e = ["维度", "指标", "维度内熵权", "维度间等权", "最终权重"]
for i, h in enumerate(headers_e):
    set_cell_font(table_e.rows[0].cells[i], h, bold=True)
for idx, (_, row) in enumerate(weight_df.iterrows()):
    r = idx + 1
    set_cell_font(table_e.rows[r].cells[0], row["dimension"], alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(table_e.rows[r].cells[1], row["indicator"], alignment=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_font(table_e.rows[r].cells[2], f"{row['local_weight']:.6f}")
    set_cell_font(table_e.rows[r].cells[3], "0.250000")
    set_cell_font(table_e.rows[r].cells[4], f"{row['weight']:.6f}")

doc.add_page_break()

# ============================================================
# 附录F：数据源与处理参数详表
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("附录F  数据源与处理参数详表")
run.font.name = "黑体"
run.font.size = Pt(14)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

p = doc.add_paragraph()
run = p.add_run("表F-1  遥感栅格数据源与处理参数")
run.font.name = "宋体"
run.font.size = Pt(9)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

gee_data = [
    ["pop_density", "CIESIN/GPWv411/GPW_Population_Density", "population_density", "2010/2015/2020/2020*", "928", "reduceResolution(mean)→10km"],
    ["nightlight", "NOAA/VIIRS/DNB/MONTHLY_V1", "avg_rad", "2012*/2015/2020/2024", "500", "年均, reduceResolution(mean)→10km"],
    ["cropland_ratio", "MODIS/061/MCD12Q1", "LC_Type1(12,14)", "2010/2015/2020/2024", "500", "像元比例, reduceResolution(mean)→10km"],
    ["forest_ratio", "MODIS/061/MCD12Q1", "LC_Type1(1-5)", "2010/2015/2020/2024", "500", "像元比例, reduceResolution(mean)→10km"],
    ["ndvi", "MODIS/061/MOD13A3", "NDVI", "2010/2015/2020/2024", "1000", "5-10月均值×0.0001, reduceResolution(mean)→10km"],
    ["npp", "MODIS/061/MOD17A3HGF", "Npp", "2010/2015/2020/2024", "500", "×0.0001, reduceResolution(mean)→10km"],
    ["slope", "USGS/GTOPO30", "elevation", "静态", "1000", "ee.Terrain.slope(), reduceResolution(mean)→10km"],
    ["climate_water_stress", "IDAHO_EPSCOR/TERRACLIMATE", "def", "2010/2015/2020/2024", "4638", "年求和×0.1, reduceResolution(mean)→10km"],
    ["cold_stress", "IDAHO_EPSCOR/TERRACLIMATE", "tmmn", "2010/2015/2020/2024", "4638", "×0.1取负max(0)后年求和, reduceResolution(mean)→10km"],
]

table_f = doc.add_table(rows=len(gee_data) + 1, cols=6)
table_f.alignment = WD_TABLE_ALIGNMENT.CENTER
make_three_line_table(table_f)
headers_f = ["指标", "GEE数据集", "波段", "实际年份", "原始分辨率(m)", "处理方式"]
for i, h in enumerate(headers_f):
    set_cell_font(table_f.rows[0].cells[i], h, bold=True)
for idx, row_data in enumerate(gee_data):
    for j, val in enumerate(row_data):
        set_cell_font(table_f.rows[idx + 1].cells[j], val, alignment=WD_ALIGN_PARAGRAPH.LEFT)

p = doc.add_paragraph()
run = p.add_run("注：*2024年人口采用2020年GPW基线；2010年夜光采用2012年VIIRS数据。所有数据在GEE端按研究区边界（外扩20km）裁剪后导出，坐标系为EPSG:4326。")
run.font.name = "宋体"
run.font.size = Pt(8)
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

doc.add_page_break()

# ============================================================
# 附录G：LISA维度短板统计
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("附录G  LISA维度短板统计")
run.font.name = "黑体"
run.font.size = Pt(14)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")

p = doc.add_paragraph()
run = p.add_run("表G-1  各LISA集聚类型的维度短板与约束指标")
run.font.name = "宋体"
run.font.size = Pt(9)
run.font.bold = True
run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

lisa_df = pd.read_csv(f"{DATA_DIR}/lisa_dimension_shortboard_statistics.csv")
table_g = doc.add_table(rows=len(lisa_df) + 1, cols=6)
table_g.alignment = WD_TABLE_ALIGNMENT.CENTER
make_three_line_table(table_g)
headers_g = ["LISA类型", "网格数", "主短板维度", "次短板维度", "主约束指标", "短板强度"]
for i, h in enumerate(headers_g):
    set_cell_font(table_g.rows[0].cells[i], h, bold=True)
for idx, (_, row) in enumerate(lisa_df.iterrows()):
    r = idx + 1
    set_cell_font(table_g.rows[r].cells[0], str(row.get("lisa_type", row.get("cluster_type", ""))))
    set_cell_font(table_g.rows[r].cells[1], str(int(row.get("grid_count", row.get("count", 0)))))
    set_cell_font(table_g.rows[r].cells[2], str(row.get("primary_shortboard", row.get("worst_dim", ""))))
    set_cell_font(table_g.rows[r].cells[3], str(row.get("secondary_shortboard", row.get("second_worst_dim", ""))))
    set_cell_font(table_g.rows[r].cells[4], str(row.get("constraint_indicator", row.get("worst_indicator", ""))))
    set_cell_font(table_g.rows[r].cells[5], f"{row.get('shortboard_intensity', row.get('worst_intensity', 0)):.4f}" if isinstance(row.get('shortboard_intensity', row.get('worst_intensity', 0)), (int, float)) else str(row.get('shortboard_intensity', row.get('worst_intensity', ""))))

doc.save(OUT_PATH)
print(f"Saved to {OUT_PATH}")
